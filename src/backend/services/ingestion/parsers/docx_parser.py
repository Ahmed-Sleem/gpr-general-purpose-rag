"""
DOCX Word Document Parser (`src/backend/services/ingestion/parsers/docx_parser.py`).

Extracts paragraphs, heading styles (`Heading 1`, `Heading 2`), and document tables using `python-docx`.
"""

import docx
from typing import List, Dict, Any
from ..normalizer import fix_kerning


def parse_docx_file(file_path: str) -> List[Dict[str, Any]]:
    """Extract structural blocks (`heading`, `text`, `table`) from a Word `.docx` document."""
    blocks: List[Dict[str, Any]] = []
    doc = docx.Document(file_path)

    # 1. Paragraphs and Headings
    for idx, para in enumerate(doc.paragraphs):
        text = fix_kerning(para.text.strip())
        if not text:
            continue

        style_name = str(para.style.name).lower() if para.style else ""
        if "heading 1" in style_name or "title" in style_name:
            blocks.append({
                "type": "heading",
                "level": 1,
                "code": f"h1_{idx}",
                "title": text,
                "content": text,
                "page_number": 1
            })
        elif "heading 2" in style_name:
            blocks.append({
                "type": "heading",
                "level": 2,
                "code": f"h2_{idx}",
                "title": text,
                "content": text,
                "page_number": 1
            })
        elif "heading 3" in style_name:
            blocks.append({
                "type": "heading",
                "level": 3,
                "code": f"h3_{idx}",
                "title": text,
                "content": text,
                "page_number": 1
            })
        else:
            if blocks and blocks[-1]["type"] == "text":
                blocks[-1]["content"] += f"\n{text}"
            else:
                blocks.append({
                    "type": "text",
                    "level": 0,
                    "title": text[:50] + ("..." if len(text) > 50 else ""),
                    "content": text,
                    "page_number": 1
                })

    # 2. DOCX Tables
    for t_idx, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            row_cells = [fix_kerning(cell.text.strip().replace("\n", " ")) for cell in row.cells if cell.text.strip()]
            if row_cells:
                rows_data.append(row_cells)

        if len(rows_data) >= 2:
            headers = rows_data[0]
            rows = rows_data[1:]
            blocks.append({
                "type": "table",
                "level": 2,
                "title": f"Table {t_idx+1}: {headers[0] if headers else 'DOCX Table'}",
                "content": f"Headers: {' | '.join(headers)}\nRows:\n" + "\n".join([" | ".join(r) for r in rows[:15]]),
                "page_number": 1,
                "table_headers": headers,
                "table_rows": rows
            })

    return blocks
